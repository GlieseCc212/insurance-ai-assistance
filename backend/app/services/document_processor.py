"""
Document Processing Service
Handles PDF and Word document text extraction and processing
"""
import os
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path
import asyncio
import time

import PyPDF2
import pdfplumber
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

from ..config import get_settings

settings = get_settings()


class DocumentProcessor:
    """Service for processing insurance policy documents"""
    
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
        # Initialize text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
    
    async def process_document(
        self, 
        file_content: bytes, 
        filename: str, 
        policy_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Process an uploaded document and extract text
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            policy_type: Type of insurance policy
            
        Returns:
            Dictionary with processing results
        """
        start_time = time.time()
        
        # Generate unique document ID
        doc_id = str(uuid.uuid4())
        
        # Get file extension
        file_extension = Path(filename).suffix.lower()
        
        # Save file temporarily
        temp_file_path = self.upload_dir / f"{doc_id}{file_extension}"
        with open(temp_file_path, "wb") as f:
            f.write(file_content)
        
        try:
            # Extract text based on file type
            if file_extension == '.pdf':
                text_content = await self._extract_pdf_text(temp_file_path)
                pages_processed = self._get_pdf_page_count(temp_file_path)
            elif file_extension in ['.docx', '.doc']:
                text_content = await self._extract_docx_text(temp_file_path)
                pages_processed = 1  # DOCX doesn't have clear page concept
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Split text into chunks
            chunks = await self._create_text_chunks(text_content, doc_id)
            
            # Extract metadata
            metadata = await self._extract_metadata(text_content, policy_type)
            
            processing_time = time.time() - start_time
            
            result = {
                'document_id': doc_id,
                'filename': filename,
                'file_type': file_extension,
                'policy_type': policy_type,
                'text_length': len(text_content),
                'pages_processed': pages_processed,
                'chunks_created': len(chunks),
                'processing_time': round(processing_time, 2),
                'metadata': metadata,
                'chunks': chunks  # This will be stored in vector DB
            }
            
            return result
            
        finally:
            # Clean up temporary file
            if temp_file_path.exists():
                temp_file_path.unlink()
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using pdfplumber (better for tables/layout)"""
        text_content = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
        except Exception as e:
            # Fallback to PyPDF2 if pdfplumber fails
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
        
        return "\n\n".join(text_content)
    
    def _get_pdf_page_count(self, file_path: Path) -> int:
        """Get the number of pages in a PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except:
            return 0
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from Word document"""
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n\n".join(text_content)
    
    async def _create_text_chunks(self, text: str, doc_id: str) -> List[Dict[str, Any]]:
        """Split text into chunks for vector storage"""
        documents = [LangChainDocument(page_content=text, metadata={"document_id": doc_id})]
        chunks = self.text_splitter.split_documents(documents)
        
        chunk_data = []
        for i, chunk in enumerate(chunks):
            chunk_data.append({
                'chunk_id': f"{doc_id}_chunk_{i}",
                'document_id': doc_id,
                'chunk_index': i,
                'content': chunk.page_content,
                'metadata': chunk.metadata
            })
        
        return chunk_data
    
    async def _extract_metadata(self, text: str, policy_type: Optional[str]) -> Dict[str, Any]:
        """Extract metadata from document text"""
        metadata = {
            'policy_type': policy_type,
            'extracted_fields': {}
        }
        
        # Simple pattern matching for common insurance fields
        patterns = {
            'policy_number': [
                r'Policy\s+Number:?\s*(\w+)',
                r'Policy\s+No\.?:?\s*(\w+)',
                r'Certificate\s+Number:?\s*(\w+)'
            ],
            'effective_date': [
                r'Effective\s+Date:?\s*(\d{1,2}/\d{1,2}/\d{4})',
                r'Policy\s+Period:?\s*(\d{1,2}/\d{1,2}/\d{4})'
            ],
            'expiration_date': [
                r'Expiration\s+Date:?\s*(\d{1,2}/\d{1,2}/\d{4})',
                r'Expires?:?\s*(\d{1,2}/\d{1,2}/\d{4})'
            ],
            'deductible': [
                r'Deductible:?\s*\$?(\d+(?:,\d{3})*(?:\.\d{2})?)',
                r'Annual\s+Deductible:?\s*\$?(\d+(?:,\d{3})*(?:\.\d{2})?)'
            ]
        }
        
        import re
        for field_name, field_patterns in patterns.items():
            for pattern in field_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    metadata['extracted_fields'][field_name] = match.group(1)
                    break
        
        return metadata
    
    def get_supported_file_types(self) -> List[str]:
        """Get list of supported file extensions"""
        return ['.pdf', '.docx', '.doc']
    
    def validate_file_type(self, filename: str) -> bool:
        """Check if file type is supported"""
        file_extension = Path(filename).suffix.lower()
        return file_extension in self.get_supported_file_types()
    
    def validate_file_size(self, file_size: int) -> bool:
        """Check if file size is within limits"""
        return file_size <= settings.max_file_size


# Global instance
document_processor = DocumentProcessor()